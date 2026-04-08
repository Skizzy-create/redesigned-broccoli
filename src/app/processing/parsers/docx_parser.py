# pyright: reportMissingImports=false

from __future__ import annotations

from pathlib import Path

from app.processing.parsers.base import ParsedDocument


class DocxParser:
    def parse(self, file_path: Path) -> ParsedDocument:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise RuntimeError("python-docx is required for DOCX parsing.") from exc

        document = DocxDocument(file_path)
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
        text = "\n\n".join(paragraphs)
        metadata = {"paragraph_count": len(paragraphs)}
        return ParsedDocument(text=text, page_count=1, metadata=metadata)
